import os
import tempfile
import streamlit as st
from langchain_community.document_loaders import PyPDFLoader, UnstructuredCSVLoader, TextLoader
from langchain_community.document_loaders import UnstructuredExcelLoader
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores.faiss import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from docx import Document
from langchain_core.documents import Document as LangchainDocument

hide_st_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            header {visibility: hidden;}
            </style>
            """
st.markdown(hide_st_style, unsafe_allow_html=True)

# Set up Gemini API key
Gemini = st.secrets["GOOGLE_API_KEY"]
llm = ChatGoogleGenerativeAI(model="gemini-1.5-pro-latest", google_api_key=Gemini)

 

def load_docx(file_path):
    doc = Document(file_path)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return [LangchainDocument(page_content=text, metadata={"source": file_path})]

@st.cache_resource(show_spinner=False)
def load_file(uploaded_file, file_name):
    temp_dir = tempfile.TemporaryDirectory()
    temp_file_path = os.path.join(temp_dir.name, uploaded_file.name)
    with open(temp_file_path, "wb") as f:
        f.write(uploaded_file.read())
    
    documents = []
    try:
        if uploaded_file.name.lower().endswith(".pdf"):
            loader = PyPDFLoader(temp_file_path)
            documents = loader.load_and_split()
        elif uploaded_file.name.lower().endswith(".csv"):
            loader = UnstructuredCSVLoader(temp_file_path)
            documents = loader.load_and_split()
        elif uploaded_file.name.lower().endswith(".docx"):
            documents = load_docx(temp_file_path)
        elif uploaded_file.name.lower().endswith((".xlsx", ".xls")):
            loader = UnstructuredExcelLoader(temp_file_path)
            documents = loader.load_and_split()
        elif uploaded_file.name.lower().endswith(".txt"):
            loader = TextLoader(temp_file_path)
            documents = loader.load_and_split()
        else:
            raise ValueError("Unsupported file type")
    except Exception as error:
        st.error(f"Error loading file: {error}")
        return None

    embeddings = HuggingFaceEmbeddings()
    vector = FAISS.from_documents(documents, embedding=embeddings)
    vector.save_local(f"{file_name}_INDEX")
    st.success("File processed and indexed successfully.")
    return documents

@st.cache_resource(show_spinner=False)
def get_summarized_response(embedding_path, file_upload_name):
    cache_key = f"embeddings_{embedding_path}_{file_upload_name}"
    if cache_key not in st.session_state:
        embeddings = HuggingFaceEmbeddings()
        vector = FAISS.load_local(embedding_path, embeddings, allow_dangerous_deserialization=True)
        db = vector.as_retriever()
   
        template = """You are a bot who responds to the User queries in a clear and professional manner.

        <context>
        {context}
        </context>

        Question: {input}
        Ensure your response is clear, concise, and directly relevant to the User's query.
        """

        prompt = ChatPromptTemplate.from_template(template=template)
        doc_chain = create_stuff_documents_chain(llm, prompt)
        chain = create_retrieval_chain(db, doc_chain)
        st.session_state[cache_key] = chain
    return st.session_state[cache_key]

def main():
    # st.title("Document Q&A")

    uploaded_file = st.file_uploader("Upload your file", type=["pdf", "csv", "docx", "xlsx", "xls", "txt"], label_visibility="collapsed")
    if uploaded_file is not None:
        st.success("File uploaded successfully.")
        file_name = uploaded_file.name.split('.')[0]

        with st.spinner("Processing"):
            documents = load_file(uploaded_file, file_name)
            if documents is not None:
                chain = get_summarized_response(f"{file_name}_INDEX", uploaded_file.name)

                user_input = st.chat_input("Enter your question here:")
                if user_input:
                    with st.spinner("Getting Response"):
                        response = chain.invoke({"input": user_input})
                        with st.expander("Response", expanded=True):
                            st.write(response['answer'])

if __name__ == "__main__":
    main()